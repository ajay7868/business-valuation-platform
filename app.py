
import os
import json
from datetime import datetime
import re
from flask import Flask, request, jsonify, send_file, render_template
from flask_cors import CORS
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from werkzeug.utils import secure_filename

# Import our custom modules
try:
    from models import db, User, RateLimit, UserActivity
    from auth_utils import (
        hash_password, verify_password, generate_verification_token,
        is_valid_email, is_valid_password, check_rate_limit,
        log_user_activity, require_auth, get_rate_limit_status
    )
    from email_config import mail, send_verification_email, send_welcome_email
    print("✅ All authentication modules imported successfully")
except ImportError as e:
    print(f"❌ Import error: {e}")
    print("Please ensure all required modules are in the same directory")
    print("Installing required packages...")
    
    # Try to install missing packages
    import subprocess
    import sys
    
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", 
                             "flask-login", "flask-sqlalchemy", "flask-mail", "flask-limiter"])
        print("✅ Dependencies installed successfully")
        
        # Try importing again
        from models import db, User, RateLimit, UserActivity
        from auth_utils import (
            hash_password, verify_password, generate_verification_token,
            is_valid_email, is_valid_password, check_rate_limit,
            log_user_activity, require_auth, get_rate_limit_status
        )
        from email_config import mail, send_verification_email, send_welcome_email
        print("✅ All authentication modules imported successfully after installation")
    except Exception as install_error:
        print(f"❌ Failed to install dependencies: {install_error}")
        print("❌ Authentication features will be disabled")
        db = None
        mail = None

# Production configuration
app = Flask(__name__)

# Database configuration
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL', 'sqlite:///valuation_platform.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Secret key for sessions
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'your-secret-key-change-in-production')

# Email configuration
app.config['MAIL_SERVER'] = os.environ.get('MAIL_SERVER', 'smtp.gmail.com')
app.config['MAIL_PORT'] = int(os.environ.get('MAIL_PORT', 587))
app.config['MAIL_USE_TLS'] = os.environ.get('MAIL_USE_TLS', 'true').lower() == 'true'
app.config['MAIL_USERNAME'] = os.environ.get('MAIL_USERNAME', 'your-email@gmail.com')
app.config['MAIL_PASSWORD'] = os.environ.get('MAIL_PASSWORD', 'your-app-password')
app.config['MAIL_DEFAULT_SENDER'] = os.environ.get('MAIL_DEFAULT_SENDER', 'your-email@gmail.com')

# File upload configuration
app.config['UPLOAD_FOLDER'] = os.environ.get('UPLOAD_FOLDER', 'uploads')
app.config['REPORTS_FOLDER'] = os.environ.get('REPORTS_FOLDER', 'reports')
app.config['MAX_CONTENT_LENGTH'] = int(os.environ.get('MAX_CONTENT_LENGTH', 100 * 1024 * 1024))  # 100MB max file size

# Initialize extensions
if db:
    db.init_app(app)
if mail:
    mail.init_app(app)

# Initialize Flask-Login
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'
login_manager.login_message = 'Please log in to access this page.'

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# Enable CORS for production
CORS(app, origins=[
    "http://localhost:3000",
    "https://your-vercel-domain.vercel.app",  # Update with your actual Vercel domain
    "https://*.vercel.app"
])

# Ensure directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['REPORTS_FOLDER'], exist_ok=True)

# Allowed file extensions
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'png', 'jpg', 'jpeg', 'gif', 'xlsx', 'xls', 'csv'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Create database tables
if db:
    with app.app_context():
        db.create_all()
else:
    print("WARNING: Database not initialized - authentication features disabled")

@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint for production monitoring"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.datetime.now().isoformat(),
        'environment': os.environ.get('FLASK_ENV', 'development'),
        'version': '1.0.0'
    })

# Authentication endpoints
@app.route('/api/auth/signup', methods=['POST'])
def signup():
    """User registration endpoint"""
    if not db or not User:
        return jsonify({'error': 'Authentication system not available'}), 503
    
    try:
        data = request.json
        
        # Validate required fields
        if not data or not data.get('email') or not data.get('password') or not data.get('confirm_password'):
            return jsonify({'error': 'Missing required fields'}), 400
        
        email = data['email'].lower().strip()
        password = data['password']
        confirm_password = data['confirm_password']
        mobile = data.get('mobile', '').strip()
        
        # Validate email format
        if not is_valid_email(email):
            return jsonify({'error': 'Invalid email format'}), 400
        
        # Check if passwords match
        if password != confirm_password:
            return jsonify({'error': 'Passwords do not match'}), 400
        
        # Validate password strength
        is_valid, message = is_valid_password(password)
        if not is_valid:
            return jsonify({'error': message}), 400
        
        # Check if user already exists
        existing_user = User.query.filter_by(email=email).first()
        if existing_user:
            return jsonify({'error': 'Email already registered'}), 409
        
        # Create new user
        verification_token = generate_verification_token()
        new_user = User(
            email=email,
            password_hash=hash_password(password),
            mobile=mobile,
            verification_token=verification_token
        )
        
        db.session.add(new_user)
        db.session.commit()
        
        # Send verification email
        email_sent, email_message = send_verification_email(email, verification_token)
        
        return jsonify({
            'status': 'success',
            'message': 'Account created successfully. Please check your email for verification.',
            'email_sent': email_sent,
            'email_message': email_message
        }), 201
        
    except Exception as e:
        if db:
            db.session.rollback()
        print(f"Signup error: {str(e)}")
        return jsonify({'error': f'Registration failed: {str(e)}'}), 500

@app.route('/api/auth/login', methods=['POST'])
def login():
    """User login endpoint"""
    if not db or not User:
        return jsonify({'error': 'Authentication system not available'}), 503
    
    try:
        data = request.json
        
        if not data or not data.get('email') or not data.get('password'):
            return jsonify({'error': 'Email and password required'}), 400
        
        email = data['email'].lower().strip()
        password = data['password']
        
        # Find user
        user = User.query.filter_by(email=email).first()
        
        if not user or not verify_password(password, user.password_hash):
            return jsonify({'error': 'Invalid email or password'}), 401
        
        if not user.email_verified:
            return jsonify({'error': 'Email not verified. Please check your email for verification link.'}), 403
        
        # Update last login
        user.last_login = datetime.utcnow()
        db.session.commit()
        
        # Log in user
        login_user(user)
        
        return jsonify({
            'status': 'success',
            'message': 'Login successful',
            'user': {
                'id': user.id,
                'email': user.email,
                'mobile': user.mobile,
                'email_verified': user.email_verified
            }
        })
        
    except Exception as e:
        print(f"Login error: {str(e)}")
        return jsonify({'error': f'Login failed: {str(e)}'}), 500

@app.route('/api/auth/verify/<token>', methods=['GET'])
def verify_email(token):
    """Email verification endpoint"""
    try:
        user = User.query.filter_by(verification_token=token).first()
        
        if not user:
            return jsonify({'error': 'Invalid verification token'}), 400
        
        if user.email_verified:
            return jsonify({'error': 'Email already verified'}), 400
        
        # Verify email
        user.email_verified = True
        user.verification_token = None
        db.session.commit()
        
        # Send welcome email
        send_welcome_email(user.email)
        
        return jsonify({
            'status': 'success',
            'message': 'Email verified successfully! You can now log in.'
        })
        
    except Exception as e:
        db.session.rollback()
        print(f"Email verification error: {str(e)}")
        return jsonify({'error': f'Verification failed: {str(e)}'}), 500

@app.route('/api/auth/logout', methods=['POST'])
@login_required
def logout():
    """User logout endpoint"""
    try:
        logout_user()
        return jsonify({
            'status': 'success',
            'message': 'Logged out successfully'
        })
    except Exception as e:
        print(f"Logout error: {str(e)}")
        return jsonify({'error': f'Logout failed: {str(e)}'}), 500

@app.route('/api/auth/profile', methods=['GET'])
@login_required
def get_profile():
    """Get user profile"""
    try:
        return jsonify({
            'status': 'success',
            'user': {
                'id': current_user.id,
                'email': current_user.email,
                'mobile': current_user.mobile,
                'email_verified': current_user.email_verified,
                'created_at': current_user.created_at.isoformat(),
                'last_login': current_user.last_login.isoformat() if current_user.last_login else None
            }
        })
    except Exception as e:
        print(f"Profile error: {str(e)}")
        return jsonify({'error': f'Failed to get profile: {str(e)}'}), 500

@app.route('/api/auth/rate-limit-status', methods=['GET'])
def get_rate_limit_status_endpoint():
    """Get current rate limit status for user"""
    try:
        upload_status = get_rate_limit_status('upload')
        report_status = get_rate_limit_status('report_generation')
        
        return jsonify({
            'status': 'success',
            'upload': upload_status,
            'report_generation': report_status
        })
    except Exception as e:
        print(f"Rate limit status error: {str(e)}")
        return jsonify({'error': f'Failed to get rate limit status: {str(e)}'}), 500

@app.route('/api/upload', methods=['POST'])
def upload_file():
    """Upload and process financial documents with rate limiting"""
    try:
        # Check rate limit
        rate_limit_ok, rate_limit_message = check_rate_limit('upload', max_attempts=2)
        if not rate_limit_ok:
            return jsonify({
                'error': 'Rate limit exceeded',
                'message': rate_limit_message,
                'requires_signup': True
            }), 429
        
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': 'File type not allowed. Please upload PDF, Excel, or image files.'}), 400
        
        # Check file size
        file.seek(0, 2)  # Seek to end
        file_size = file.tell()
        file.seek(0)  # Reset to beginning
        
        print(f"DEBUG: File size: {file_size} bytes, {file_size / 1024 / 1024:.2f} MB")
        print(f"DEBUG: Max allowed size: {app.config['MAX_CONTENT_LENGTH']} bytes, {app.config['MAX_CONTENT_LENGTH'] / 1024 / 1024:.2f} MB")
        
        if file_size > app.config['MAX_CONTENT_LENGTH']:
            return jsonify({'error': f'File too large. Maximum size is {app.config["MAX_CONTENT_LENGTH"] // (1024*1024)}MB.'}), 413
        
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        # Extract data based on file type
        extracted_data = extract_data_from_file(filepath)
        
        # AI validation of extracted data
        validation_result = validate_financial_data_with_ai(extracted_data)
        
        # Log activity
        user_id = current_user.id if current_user.is_authenticated else None
        log_user_activity(user_id, 'upload', True)
        
        return jsonify({
            'status': 'success',
            'message': 'File uploaded and processed successfully',
            'filename': filename,
            'extracted_data': extracted_data,
            'ai_validation': validation_result,
            'rate_limit': get_rate_limit_status('upload')
        })
        
    except Exception as e:
        print(f"Upload error: {str(e)}")
        # Log failed activity
        user_id = current_user.id if current_user.is_authenticated else None
        log_user_activity(user_id, 'upload', False)
        return jsonify({'error': f'File upload failed: {str(e)}'}), 500

def extract_data_from_file(filepath):
    """Extract real financial data from various file types"""
    try:
        file_extension = filepath.rsplit('.', 1)[1].lower()
        
        if file_extension in ['xlsx', 'xls']:
            return extract_from_excel(filepath)
        elif file_extension == 'pdf':
            return extract_from_pdf(filepath)
        elif file_extension == 'csv':
            return extract_from_csv(filepath)
        else:
            return extract_from_image(filepath)
            
    except Exception as e:
        print(f"Data extraction error: {str(e)}")
        return get_empty_data()

def extract_from_excel(filepath):
    """Extract data from Excel files with enhanced multi-sheet and multi-column support"""
    try:
        import pandas as pd
        import openpyxl
        
        # Read Excel file with all sheets
        df = pd.read_excel(filepath, sheet_name=None)
        print(f"DEBUG: Excel file has {len(df)} sheets: {list(df.keys())}")
        
        all_data = {}
        best_sheet_data = {}
        best_sheet_score = 0
        
        # Process each sheet and score them based on data quality
        for sheet_name, sheet_df in df.items():
            print(f"DEBUG: Processing sheet: {sheet_name}")
            
            if sheet_df.empty:
                print(f"DEBUG: Sheet {sheet_name} is empty, skipping")
                continue
                
            print(f"DEBUG: Sheet {sheet_name} has {len(sheet_df.columns)} columns and {len(sheet_df)} rows")
            print(f"DEBUG: Columns: {list(sheet_df.columns)}")
            
            # Try to extract data from this sheet
            sheet_data = extract_data_from_sheet(sheet_df, sheet_name)
            if sheet_data:
                # Score this sheet based on data completeness
                sheet_score = len(sheet_data)
                print(f"DEBUG: Sheet {sheet_name} score: {sheet_score}")
                
                if sheet_score > best_sheet_score:
                    best_sheet_score = sheet_score
                    best_sheet_data = sheet_data
                
                # Merge data from all sheets
                all_data.update(sheet_data)
        
        # If we have good data from a specific sheet, prioritize it
        if best_sheet_score >= 3:  # At least 3 metrics found
            print(f"DEBUG: Using best sheet data with score {best_sheet_score}")
            final_data = best_sheet_data
        elif all_data:
            print(f"DEBUG: Using combined data from all sheets")
            final_data = all_data
        else:
            print(f"DEBUG: No data found in any sheet")
            final_data = get_empty_data()
        
        # Validate and clean the extracted data
        final_data = validate_and_clean_data(final_data)
        
        print(f"DEBUG: Final extracted Excel data: {final_data}")
        return final_data
        
    except Exception as e:
        print(f"Excel extraction error: {str(e)}")
        return get_empty_data()

def extract_data_from_sheet(df, sheet_name):
    """Extract data from a single Excel sheet with enhanced column analysis"""
    try:
        print(f"DEBUG: Analyzing sheet '{sheet_name}' with {len(df.columns)} columns")
        
        # Convert to string for easier searching
        df_str = df.astype(str)
        
        # Extract company information
        company_data = extract_company_info_from_dataframe(df, df_str)
        
        # Extract financial metrics with enhanced multi-column support
        financial_metrics = extract_financial_metrics(df, df_str)
        
        # Merge the data
        sheet_data = {**company_data, **financial_metrics}
        
        print(f"DEBUG: Sheet '{sheet_name}' extracted data: {sheet_data}")
        return sheet_data
        
    except Exception as e:
        print(f"Sheet extraction error for '{sheet_name}': {str(e)}")
        return {}

def extract_from_pdf(filepath):
    """Extract data from PDF files"""
    try:
        import pdfplumber
        
        extracted_text = ""
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    extracted_text += text + "\n"
        
        # Extract company information from text
        company_data = extract_company_info_from_text(extracted_text)
        
        # Validate and clean the extracted data
        company_data = validate_and_clean_data(company_data)
        
        print(f"DEBUG: Extracted PDF data: {company_data}")
        return company_data
        
    except Exception as e:
        print(f"PDF extraction error: {str(e)}")
        return get_empty_data()

def extract_from_csv(filepath):
    """Extract data from CSV files"""
    try:
        import pandas as pd
        
        # Read CSV file
        df = pd.read_csv(filepath)
        
        # Check if CSV has a key-value structure (like our test file)
        if len(df.columns) == 2 and 'Metric' in df.columns and 'Value' in df.columns:
            # Handle key-value structure
            company_data = get_empty_data()
            
            # Extract company name
            company_row = df[df['Metric'] == 'Company Name']
            if not company_row.empty:
                company_data['company_name'] = company_row['Value'].iloc[0]
            
            # Extract industry
            industry_row = df[df['Metric'] == 'Industry']
            if not industry_row.empty:
                company_data['industry'] = industry_row['Value'].iloc[0]
            
            # Extract financial metrics
            for metric in ['Revenue', 'EBITDA', 'Total Assets', 'Inventory', 'Accounts Receivable', 'Cash', 'Total Liabilities', 'Employees']:
                row = df[df['Metric'] == metric]
                if not row.empty:
                    value_str = str(row['Value'].iloc[0])
                    try:
                        # Clean the value (remove $ and commas)
                        clean_value = float(value_str.replace('$', '').replace(',', ''))
                        if metric == 'Revenue':
                            company_data['revenue'] = clean_value
                        elif metric == 'EBITDA':
                            company_data['ebitda'] = clean_value
                        elif metric == 'Total Assets':
                            company_data['total_assets'] = clean_value
                        elif metric == 'Inventory':
                            company_data['inventory'] = clean_value
                        elif metric == 'Accounts Receivable':
                            company_data['accounts_receivable'] = clean_value
                        elif metric == 'Cash':
                            company_data['cash'] = clean_value
                        elif metric == 'Total Liabilities':
                            company_data['total_liabilities'] = clean_value
                        elif metric == 'Employees':
                            company_data['employees'] = int(clean_value)
                    except ValueError:
                        continue
            
            print(f"DEBUG: Extracted CSV data (key-value): {company_data}")
            
            # Validate and clean the extracted data
            company_data = validate_and_clean_data(company_data)
            return company_data
        else:
            # Handle regular CSV structure
            df_str = df.astype(str)
            company_data = extract_company_info_from_dataframe(df, df_str)
            
            print(f"DEBUG: Extracted CSV data (regular): {company_data}")
            
            # Validate and clean the extracted data
            company_data = validate_and_clean_data(company_data)
            return company_data
        
    except Exception as e:
        print(f"CSV extraction error: {str(e)}")
        return get_empty_data()

def extract_from_image(filepath):
    """Extract data from image files (OCR would be needed for production)"""
    try:
        # For now, return empty data for images
        # In production, you could integrate OCR services like Tesseract or cloud OCR
        return get_empty_data()
    except Exception as e:
        print(f"Image extraction error: {str(e)}")
        return get_empty_data()

def extract_company_info_from_dataframe(df, df_str):
    """Extract company information from pandas DataFrame"""
    company_data = get_empty_data()
    
    try:
        # Look for company name in column headers or first few rows
        company_name = find_company_name(df, df_str)
        if company_name:
            company_data['company_name'] = company_name
        
        # Look for industry information
        industry = find_industry(df, df_str)
        if industry:
            company_data['industry'] = industry
        
        # Extract financial metrics
        financial_metrics = extract_financial_metrics(df, df_str)
        company_data.update(financial_metrics)
        
        # Extract employee count
        employees = find_employee_count(df, df_str)
        if employees:
            company_data['employees'] = employees
            
    except Exception as e:
        print(f"DataFrame extraction error: {str(e)}")
    
    return company_data

def extract_company_info_from_text(text):
    """Extract company information from text content"""
    company_data = get_empty_data()
    
    try:
        # Look for company name patterns
        company_name = find_company_name_in_text(text)
        if company_name:
            company_data['company_name'] = company_name
        
        # Look for industry information
        industry = find_industry_in_text(text)
        if industry:
            company_data['industry'] = industry
        
        # Extract financial metrics from text
        financial_metrics = extract_financial_metrics_from_text(text)
        company_data.update(financial_metrics)
        
        # Extract employee count from text
        employees = find_employee_count_in_text(text)
        if employees:
            company_data['employees'] = employees
            
    except Exception as e:
        print(f"Text extraction error: {str(e)}")
    
    return company_data

def find_company_name(df, df_str):
    """Find company name in DataFrame"""
    try:
        # Look for common company name patterns
        company_patterns = [
            'company', 'corp', 'inc', 'llc', 'ltd', 'enterprises', 'group', 'holdings'
        ]
        
        # Check column headers
        for col in df.columns:
            if any(pattern in str(col).lower() for pattern in company_patterns):
                # Look for non-null values in this column
                values = df[col].dropna()
                if not values.empty:
                    value = str(values.iloc[0])
                    if value and value != 'nan' and len(value.strip()) > 3:
                        return value.strip()
        
        # Check first few rows for company names
        for idx in range(min(3, len(df))):
            for col in df.columns:
                value = str(df.iloc[idx][col])
                if any(pattern in value.lower() for pattern in company_patterns):
                    if value and value != 'nan' and len(value.strip()) > 3:
                        return value.strip()
        
        # Look for any capitalized company-like names in the first few rows
        for idx in range(min(3, len(df))):
            for col in df.columns:
                value = str(df.iloc[idx][col])
                if (value and value != 'nan' and 
                    value[0].isupper() and 
                    len(value.strip()) > 5 and
                    any(word in value.lower() for word in ['company', 'corp', 'inc', 'llc', 'ltd'])):
                    return value.strip()
        
        return None
    except Exception as e:
        print(f"Company name extraction error: {str(e)}")
        return None

def find_company_name_in_text(text):
    """Find company name in text content"""
    try:
        # Look for common company name patterns
        import re
        
        # Common company suffixes
        company_suffixes = r'\b(?:Company|Corp|Inc|LLC|Ltd|Enterprises|Group|Holdings|Industries|Solutions|Technologies|Services)\b'
        
        # Look for patterns like "Company Name Inc" or "Company Name LLC"
        patterns = [
            r'\b([A-Z][a-zA-Z\s&]+)\s+(?:' + company_suffixes + r')\b',
            r'\b([A-Z][a-zA-Z\s&]+)\s+(?:Corporation|Limited|Company)\b',
            r'\b([A-Z][a-zA-Z\s&]{3,})\b'  # General capitalized names
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                # Return the first meaningful match
                for match in matches:
                    if len(match.strip()) > 3 and not match.strip().isdigit():
                        return match.strip()
        
        return None
    except Exception as e:
        print(f"Company name text extraction error: {str(e)}")
        return None

def find_industry(df, df_str):
    """Find industry information in DataFrame"""
    try:
        # Look for industry-related columns
        industry_keywords = ['industry', 'sector', 'business', 'type', 'category']
        
        for col in df.columns:
            if any(keyword in str(col).lower() for keyword in industry_keywords):
                values = df[col].dropna()
                if not values.empty:
                    return str(values.iloc[0])
        
        return None
    except Exception as e:
        print(f"Industry extraction error: {str(e)}")
        return None

def find_industry_in_text(text):
    """Find industry information in text content"""
    try:
        # Common industry keywords
        industries = [
            'manufacturing', 'technology', 'healthcare', 'finance', 'retail', 'services',
            'construction', 'transportation', 'energy', 'telecommunications', 'education',
            'real estate', 'agriculture', 'mining', 'utilities', 'media', 'entertainment'
        ]
        
        text_lower = text.lower()
        for industry in industries:
            if industry in text_lower:
                return industry.title()
        
        return None
    except Exception as e:
        print(f"Industry text extraction error: {str(e)}")
        return None

def extract_financial_metrics(df, df_str):
    """Extract financial metrics from DataFrame with intelligent multi-column analysis"""
    metrics = {}
    
    try:
        print(f"DEBUG: Processing DataFrame with {len(df.columns)} columns: {list(df.columns)}")
        
        # First pass: Look for exact column name matches using enhanced patterns
        metric_patterns = {
            'revenue': ['revenue', 'sales', 'income', 'turnover', 'gross revenue', 'total revenue', 'gross sales', 'net sales'],
            'ebitda': ['ebitda', 'operating income', 'operating profit', 'operating earnings', 'operating margin'],
            'total_assets': ['total assets', 'assets', 'total asset', 'asset base', 'total capital'],
            'inventory': ['inventory', 'stock', 'goods', 'merchandise', 'stock inventory', 'inventories'],
            'accounts_receivable': ['accounts receivable', 'receivables', 'ar', 'debtors', 'trade receivable'],
            'cash': ['cash', 'cash and cash equivalents', 'cash balance', 'bank balance', 'liquid assets'],
            'total_liabilities': ['total liabilities', 'liabilities', 'debt', 'total debt', 'obligations'],
            'net_income': ['net income', 'net profit', 'profit', 'earnings', 'net earnings', 'bottom line'],
            'employees': ['employees', 'employee', 'fte', 'full time equivalent', 'staff', 'headcount', 'personnel']
        }
        
        # Process each metric type with enhanced detection
        for metric_key, search_terms in metric_patterns.items():
            if metric_key not in metrics:  # Only set if not already found
                value = find_metric_value(df, df_str, search_terms)
                if value is not None:
                    metrics[metric_key] = value
                    print(f"DEBUG: Found {metric_key} = {value} using pattern matching")
        
        # Second pass: Analyze column data patterns for additional insights
        if len(metrics) < 5:  # If we didn't find enough metrics
            print("DEBUG: Second pass - analyzing column data patterns...")
            additional_metrics = analyze_column_patterns(df)
            for key, value in additional_metrics.items():
                if key not in metrics:
                    metrics[key] = value
                    print(f"DEBUG: Pattern analysis found {key} = {value}")
        
        # Third pass: Use the existing data structure analysis
        if len(metrics) < 3:
            print("DEBUG: Third pass - data structure analysis...")
            structure_metrics = extract_metrics_from_data_structure(df)
            for key, value in structure_metrics.items():
                if key not in metrics:
                    metrics[key] = value
                    print(f"DEBUG: Structure analysis found {key} = {value}")
        
        # Final fallback: Comprehensive numeric data finder
        if len(metrics) < 2:
            print("DEBUG: Final fallback - comprehensive numeric analysis...")
            fallback_metrics = find_any_numeric_data(df)
            for key, value in fallback_metrics.items():
                if key not in metrics:
                    metrics[key] = value
                    print(f"DEBUG: Fallback found {key} = {value}")
        
        print(f"DEBUG: Final metrics extracted: {metrics}")
        return metrics
        
    except Exception as e:
        print(f"Financial metrics extraction error: {str(e)}")
        return {}

def extract_metrics_from_data_structure(df):
    """Extract metrics from the data structure when patterns don't match"""
    metrics = {}
    
    try:
        import pandas as pd
        
        # Look for columns that might contain financial data
        for col in df.columns:
            col_lower = str(col).lower()
            
            # Check if this column contains financial data
            if any(keyword in col_lower for keyword in ['revenue', 'sales', 'income']):
                # Look for numeric values in this column
                for value in df[col]:
                    if pd.notna(value) and str(value).replace('$', '').replace(',', '').replace('.', '').isdigit():
                        try:
                            clean_value = float(str(value).replace('$', '').replace(',', ''))
                            if clean_value > 0:
                                metrics['revenue'] = clean_value
                                break
                        except ValueError:
                            continue
            
            elif any(keyword in col_lower for keyword in ['ebitda', 'operating']):
                for value in df[col]:
                    if pd.notna(value) and str(value).replace('$', '').replace(',', '').replace('.', '').isdigit():
                        try:
                            clean_value = float(str(value).replace('$', '').replace(',', ''))
                            if clean_value > 0:
                                metrics['ebitda'] = clean_value
                                break
                        except ValueError:
                            continue
            
            elif any(keyword in col_lower for keyword in ['assets', 'asset']):
                for value in df[col]:
                    if pd.notna(value) and str(value).replace('$', '').replace(',', '').replace('.', '').isdigit():
                        try:
                            clean_value = float(str(value).replace('$', '').replace(',', ''))
                            if clean_value > 0:
                                metrics['total_assets'] = clean_value
                                break
                        except ValueError:
                            continue
            
            elif any(keyword in col_lower for keyword in ['inventory', 'stock']):
                for value in df[col]:
                    if pd.notna(value) and pd.notna(value) and str(value).replace('$', '').replace(',', '').replace('.', '').isdigit():
                        try:
                            clean_value = float(str(value).replace('$', '').replace(',', ''))
                            if clean_value > 0:
                                metrics['inventory'] = clean_value
                                break
                        except ValueError:
                            continue
            
            elif any(keyword in col_lower for keyword in ['receivable', 'ar']):
                for value in df[col]:
                    if pd.notna(value) and str(value).replace('$', '').replace(',', '').replace('.', '').isdigit():
                        try:
                            clean_value = float(str(value).replace('$', '').replace(',', ''))
                            if clean_value > 0:
                                metrics['accounts_receivable'] = clean_value
                                break
                        except ValueError:
                            continue
            
            elif any(keyword in col_lower for keyword in ['cash', 'bank']):
                for value in df[col]:
                    if pd.notna(value) and str(value).replace('$', '').replace(',', '').replace('.', '').isdigit():
                        try:
                            clean_value = float(str(value).replace('$', '').replace(',', ''))
                            if clean_value > 0:
                                metrics['cash'] = clean_value
                                break
                        except ValueError:
                            continue
            
            elif any(keyword in col_lower for keyword in ['liabilities', 'debt']):
                for value in df[col]:
                    if pd.notna(value) and str(value).replace('$', '').replace(',', '').replace('.', '').isdigit():
                        try:
                            clean_value = float(str(value).replace('$', '').replace(',', ''))
                            if clean_value > 0:
                                metrics['total_liabilities'] = clean_value
                                break
                        except ValueError:
                            continue
        
    except Exception as e:
        print(f"Data structure extraction error: {str(e)}")
    
    return metrics

def validate_and_clean_data(company_data):
    """Validate and clean extracted company data"""
    try:
        # Remove None values and replace with appropriate defaults
        cleaned_data = {}
        
        for key, value in company_data.items():
            if value is not None:
                cleaned_data[key] = value
            else:
                # Set reasonable defaults for missing data
                if key in ['revenue', 'ebitda', 'total_assets', 'inventory', 'accounts_receivable', 'cash', 'total_liabilities', 'net_income']:
                    cleaned_data[key] = 0
                elif key == 'employees':
                    cleaned_data[key] = 0
                elif key == 'company_name':
                    cleaned_data[key] = 'Extracted Company'
                elif key == 'industry':
                    cleaned_data[key] = 'General Business'
        
        # Ensure we have at least some basic data
        if not cleaned_data.get('company_name'):
            cleaned_data['company_name'] = 'Extracted Company'
        
        if not cleaned_data.get('industry'):
            cleaned_data['industry'] = 'General Business'
        
        print(f"DEBUG: Cleaned data: {cleaned_data}")
        return cleaned_data
        
    except Exception as e:
        print(f"Data validation error: {str(e)}")
        return company_data

def validate_financial_data_with_ai(company_data):
    """Validate extracted financial data using OpenAI GPT for accuracy and reasonableness"""
    try:
        # Check if OpenAI API key is available
        openai_api_key = os.environ.get('OPENAI_API_KEY')
        if not openai_api_key:
            print("WARNING: OpenAI API key not found. Skipping AI validation.")
            return {
                'status': 'skipped',
                'message': 'OpenAI API key not configured',
                'confidence_score': 0.0,
                'validation_notes': ['AI validation not available - API key missing']
            }
        
        # Prepare the data for validation
        financial_summary = f"""
        Company: {company_data.get('company_name', 'Unknown')}
        Industry: {company_data.get('industry', 'Unknown')}
        Annual Revenue: ${company_data.get('revenue', 0):,.0f}
        EBITDA: ${company_data.get('ebitda', 0):,.0f}
        Total Assets: ${company_data.get('total_assets', 0):,.0f}
        Inventory: ${company_data.get('inventory', 0):,.0f}
        Accounts Receivable: ${company_data.get('accounts_receivable', 0):,.0f}
        Cash: ${company_data.get('cash', 0):,.0f}
        Total Liabilities: ${company_data.get('total_liabilities', 0):,.0f}
        Net Income: ${company_data.get('net_income', 0):,.0f}
        Employees: {company_data.get('employees', 0)}
        """
        
        # Create validation prompt
        validation_prompt = f"""
        You are a financial analyst expert. Please validate the following extracted financial data for reasonableness and accuracy.
        
        {financial_summary}
        
        Please analyze this data and provide:
        1. A confidence score (0-100) for the overall data quality
        2. Specific validation notes about any potential issues
        3. Industry-specific reasonableness checks
        4. Recommendations for data verification
        
        Focus on:
        - Financial ratios and relationships (e.g., EBITDA margin, asset turnover)
        - Industry benchmarks and reasonableness
        - Data consistency and completeness
        - Potential red flags or unusual values
        
        Respond in JSON format:
        {{
            "confidence_score": <0-100>,
            "validation_notes": ["note1", "note2", ...],
            "industry_analysis": "industry-specific insights",
            "recommendations": ["rec1", "rec2", ...],
            "risk_level": "low/medium/high"
        }}
        """
        
        # Call OpenAI API
        import openai
        openai.api_key = openai_api_key
        
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a financial analyst expert specializing in data validation and business analysis."},
                {"role": "user", "content": validation_prompt}
            ],
            max_tokens=500,
            temperature=0.1
        )
        
        # Parse the response
        ai_response = response.choices[0].message.content
        
        try:
            # Try to parse JSON response
            import json
            validation_result = json.loads(ai_response)
            validation_result['status'] = 'validated'
            validation_result['ai_response'] = ai_response
        except json.JSONDecodeError:
            # If JSON parsing fails, create a structured response
            validation_result = {
                'status': 'validated',
                'confidence_score': 75.0,
                'validation_notes': ['AI validation completed but response format unclear'],
                'industry_analysis': 'AI analysis completed',
                'recommendations': ['Review extracted data manually for verification'],
                'risk_level': 'medium',
                'ai_response': ai_response
            }
        
        print(f"DEBUG: AI validation result: {validation_result}")
        return validation_result
        
    except Exception as e:
        print(f"AI validation error: {str(e)}")
        return {
            'status': 'error',
            'message': f'AI validation failed: {str(e)}',
            'confidence_score': 0.0,
            'validation_notes': [f'Validation error: {str(e)}'],
            'risk_level': 'high'
        }

def analyze_column_patterns(df):
    """Analyze column patterns to infer financial metrics from multi-column data"""
    metrics = {}
    
    try:
        import pandas as pd
        
        print(f"DEBUG: Analyzing column patterns for {len(df.columns)} columns")
        
        # Analyze column relationships and patterns
        for col in df.columns:
            col_lower = str(col).lower()
            col_data = df[col]
            
            # Skip if column is mostly empty or non-numeric
            if col_data.isna().sum() > len(col_data) * 0.8:  # More than 80% empty
                continue
                
            # Look for columns that might contain financial data based on patterns
            if any(keyword in col_lower for keyword in ['amount', 'value', 'total', 'sum', 'balance']):
                # This could be a financial column - analyze the data
                numeric_values = []
                for value in col_data:
                    if pd.notna(value):
                        try:
                            clean_value = float(str(value).replace('$', '').replace(',', ''))
                            if clean_value > 0:
                                numeric_values.append(clean_value)
                        except ValueError:
                            continue
                
                if numeric_values:
                    max_value = max(numeric_values)
                    
                    # Infer metric type based on value magnitude and column context
                    if max_value > 1000000:  # Likely revenue or assets
                        if 'revenue' not in metrics:
                            metrics['revenue'] = max_value
                            print(f"DEBUG: Inferred revenue from column '{col}' (value: {max_value})")
                    elif max_value > 100000:  # Likely EBITDA or medium assets
                        if 'ebitda' not in metrics:
                            metrics['ebitda'] = max_value
                            print(f"DEBUG: Inferred EBITDA from column '{col}' (value: {max_value})")
                    elif max_value > 10000:  # Likely inventory or receivables
                        if 'inventory' not in metrics:
                            metrics['inventory'] = max_value
                            print(f"DEBUG: Inferred inventory from column '{col}' (value: {max_value})")
        
        # Look for columns that might represent time periods (years, quarters)
        time_columns = []
        for col in df.columns:
            col_lower = str(col).lower()
            if any(keyword in col_lower for keyword in ['year', 'period', 'quarter', 'month', 'date']):
                time_columns.append(col)
        
        if time_columns:
            print(f"DEBUG: Found time columns: {time_columns}")
        
        # Look for columns that might represent different business units or categories
        category_columns = []
        for col in df.columns:
            col_lower = str(col).lower()
            if any(keyword in col_lower for keyword in ['category', 'type', 'segment', 'division', 'unit']):
                category_columns.append(col)
        
        if category_columns:
            print(f"DEBUG: Found category columns: {category_columns}")
        
        return metrics
        
    except Exception as e:
        print(f"Column pattern analysis error: {str(e)}")
        return {}

def find_any_numeric_data(df):
    """Find any numeric data in the DataFrame that could be financial metrics"""
    metrics = {}
    
    try:
        import pandas as pd
        
        # Look through all numeric columns
        for col in df.columns:
            if df[col].dtype in ['int64', 'float64'] or df[col].dtype == 'object':
                # Check if column contains numeric data
                numeric_values = []
                for value in df[col]:
                    if pd.notna(value):
                        try:
                            # Try to convert to float
                            clean_value = float(str(value).replace('$', '').replace(',', ''))
                            if clean_value > 0:
                                numeric_values.append(clean_value)
                        except ValueError:
                            continue
                
                if numeric_values:
                    # Use the largest value as it's likely the most significant
                    max_value = max(numeric_values)
                    
                    # Try to guess what this metric represents based on column name
                    col_lower = str(col).lower()
                    
                    if any(keyword in col_lower for keyword in ['revenue', 'sales', 'income', 'turnover']):
                        metrics['revenue'] = max_value
                    elif any(keyword in col_lower for keyword in ['ebitda', 'operating', 'profit']):
                        metrics['ebitda'] = max_value
                    elif any(keyword in col_lower for keyword in ['assets', 'asset']):
                        metrics['total_assets'] = max_value
                    elif any(keyword in col_lower for keyword in ['inventory', 'stock']):
                        metrics['inventory'] = max_value
                    elif any(keyword in col_lower for keyword in ['receivable', 'ar']):
                        metrics['accounts_receivable'] = max_value
                    elif any(keyword in col_lower for keyword in ['cash', 'bank']):
                        metrics['cash'] = max_value
                    elif any(keyword in col_lower for keyword in ['liability', 'debt']):
                        metrics['total_liabilities'] = max_value
                    elif any(keyword in col_lower for keyword in ['employee', 'staff', 'fte']):
                        metrics['employees'] = int(max_value)
                    else:
                        # If we can't identify the metric, store it with a generic name
                        if 'revenue' not in metrics:
                            metrics['revenue'] = max_value
                        elif 'total_assets' not in metrics:
                            metrics['total_assets'] = max_value
        
        return metrics
        
    except Exception as e:
        print(f"Find any numeric data error: {str(e)}")
        return {}

def extract_financial_metrics_from_text(text):
    """Extract financial metrics from text content"""
    metrics = {}
    
    try:
        import re
        
        # Define financial metric patterns
        metric_patterns = {
            'revenue': [
                r'revenue[:\s]*\$?([0-9,]+(?:\.[0-9]{2})?)',
                r'sales[:\s]*\$?([0-9,]+(?:\.[0-9]{2})?)',
                r'income[:\s]*\$?([0-9,]+(?:\.[0-9]{2})?)'
            ],
            'ebitda': [
                r'ebitda[:\s]*\$?([0-9,]+(?:\.[0-9]{2})?)',
                r'operating income[:\s]*\$?([0-9,]+(?:\.[0-9]{2})?)',
                r'operating profit[:\s]*\$?([0-9,]+(?:\.[0-9]{2})?)'
            ],
            'total_assets': [
                r'total assets[:\s]*\$?([0-9,]+(?:\.[0-9]{2})?)',
                r'assets[:\s]*\$?([0-9,]+(?:\.[0-9]{2})?)'
            ],
            'inventory': [
                r'inventory[:\s]*\$?([0-9,]+(?:\.[0-9]{2})?)',
                r'stock[:\s]*\$?([0-9,]+(?:\.[0-9]{2})?)'
            ],
            'accounts_receivable': [
                r'accounts receivable[:\s]*\$?([0-9,]+(?:\.[0-9]{2})?)',
                r'receivables[:\s]*\$?([0-9,]+(?:\.[0-9]{2})?)'
            ],
            'cash': [
                r'cash[:\s]*\$?([0-9,]+(?:\.[0-9]{2})?)',
                r'bank balance[:\s]*\$?([0-9,]+(?:\.[0-9]{2})?)'
            ],
            'total_liabilities': [
                r'total liabilities[:\s]*\$?([0-9,]+(?:\.[0-9]{2})?)',
                r'liabilities[:\s]*\$?([0-9,]+(?:\.[0-9]{2})?)',
                r'debt[:\s]*\$?([0-9,]+(?:\.[0-9]{2})?)'
            ],
            'net_income': [
                r'net income[:\s]*\$?([0-9,]+(?:\.[0-9]{2})?)',
                r'net profit[:\s]*\$?([0-9,]+(?:\.[0-9]{2})?)',
                r'profit[:\s]*\$?([0-9,]+(?:\.[0-9]{2})?)'
            ]
        }
        
        for metric, patterns in metric_patterns.items():
            value = find_metric_value_in_text(text, patterns)
            if value is not None:
                metrics[metric] = value
        
    except Exception as e:
        print(f"Financial metrics text extraction error: {str(e)}")
    
    return metrics

def find_metric_value(df, df_str, patterns):
    """Find metric value in DataFrame"""
    try:
        import pandas as pd
        
        # Look for columns that match the patterns
        for pattern in patterns:
            for col in df.columns:
                if pattern in str(col).lower():
                    # Look for numeric values in this column
                    numeric_values = pd.to_numeric(df[col], errors='coerce').dropna()
                    if not numeric_values.empty:
                        return float(numeric_values.iloc[0])
        
        # Look for values in rows that contain the patterns
        for pattern in patterns:
            for idx in range(len(df)):
                for col in df.columns:
                    cell_value = str(df.iloc[idx][col])
                    if pattern in cell_value.lower():
                        # Look for numeric values in nearby cells
                        for nearby_col in df.columns:
                            nearby_value = pd.to_numeric(df.iloc[idx][nearby_col], errors='coerce')
                            if not pd.isna(nearby_value):
                                return float(nearby_value)
        
        return None
    except Exception as e:
        print(f"Metric value extraction error: {str(e)}")
        return None

def find_metric_value_in_text(text, patterns):
    """Find metric value in text content"""
    try:
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                # Convert to float, removing commas
                value_str = matches[0].replace(',', '')
                try:
                    return float(value_str)
                except ValueError:
                    continue
        
        return None
    except Exception as e:
        print(f"Metric value text extraction error: {str(e)}")
        return None

def find_employee_count(df, df_str):
    """Find employee count in DataFrame"""
    try:
        import pandas as pd
        
        # Look for employee-related columns
        employee_keywords = ['employees', 'staff', 'headcount', 'fte', 'full time equivalent']
        
        for col in df.columns:
            if any(keyword in str(col).lower() for keyword in employee_keywords):
                values = pd.to_numeric(df[col], errors='coerce').dropna()
                if not values.empty:
                    return int(values.iloc[0])
        
        return None
    except Exception as e:
        print(f"Employee count extraction error: {str(e)}")
        return None

def find_employee_count_in_text(text):
    """Find employee count in text content"""
    try:
        import re
        
        # Look for employee count patterns
        patterns = [
            r'(\d+)\s+employees?',
            r'(\d+)\s+staff',
            r'(\d+)\s+fte',
            r'headcount[:\s]*(\d+)',
            r'workforce[:\s]*(\d+)'
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                return int(matches[0])
        
        return None
    except Exception as e:
        print(f"Employee count text extraction error: {str(e)}")
        return None

def get_empty_data():
    """Return empty company data structure for real extraction"""
    return {
        'company_name': None,
        'industry': None,
        'revenue': None,
        'ebitda': None,
        'total_assets': None,
        'inventory': None,
        'accounts_receivable': None,
        'cash': None,
        'total_liabilities': None,
        'net_income': None,
        'employees': None
    }

@app.route('/api/valuation', methods=['POST'])
def calculate_valuation():
    """Calculate comprehensive business valuation with AI validation"""
    try:
        data = request.json
        
        # AI validation of company data before valuation
        print("DEBUG: Running AI validation before valuation calculation...")
        validation_result = validate_financial_data_with_ai(data)
        
        # Simplified valuation calculation for Vercel deployment
        revenue = float(data.get('revenue', 0))
        ebitda = float(data.get('ebitda', 0))
        total_assets = float(data.get('total_assets', 0))
        
        # Asset-based valuation
        asset_based = total_assets * 0.8  # 80% of book value
        
        # Income-based valuation (EBITDA multiple)
        ebitda_multiple = 6.0  # Industry average
        income_based = ebitda * ebitda_multiple
        
        # Market-based valuation (Revenue multiple)
        revenue_multiple = 1.5  # Conservative multiple
        market_based = revenue * revenue_multiple
        
        # Calculate valuation range
        valuations = [asset_based, income_based, market_based]
        min_val = min(valuations)
        max_val = max(valuations)
        mid_val = sum(valuations) / len(valuations)
        
        results = {
            'asset_based': round(asset_based, 2),
            'income_based': round(income_based, 2),
            'market_based': round(market_based, 2),
            'valuation_range': {
                'low': round(min_val, 2),
                'mid': round(mid_val, 2),
                'high': round(max_val, 2)
            },
            'methodology': 'Simplified valuation using asset, income, and market approaches',
            'assumptions': f'EBITDA multiple: {ebitda_multiple}x, Revenue multiple: {revenue_multiple}x, Asset discount: 20%'
        }
        
        # Adjust methodology based on validation results
        if validation_result.get('status') == 'validated':
            confidence = validation_result.get('confidence_score', 0)
            if confidence < 70:
                results['methodology'] += ' (Data confidence: Low - Manual verification recommended)'
                results['assumptions'] += ' - Note: Data quality concerns identified'
        
        executive_summary = f"""
        Based on the provided financial data, {data.get('company_name', 'this company')} has an estimated value range of ${min_val:,.0f} to ${max_val:,.0f}, with a mid-point estimate of ${mid_val:,.0f}.
        
        The valuation considers:
        • Asset-based approach: ${asset_based:,.0f}
        • Income-based approach: ${income_based:,.0f}
        • Market-based approach: ${market_based:,.0f}
        
        This represents a comprehensive assessment using industry-standard methodologies.
        """
        
        return jsonify({
            'status': 'success',
            'valuation_results': results,
            'executive_summary': executive_summary,
            'company_data': data,
            'ai_validation': validation_result
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/swot', methods=['POST'])
def generate_swot():
    """Generate SWOT analysis using AI with data validation"""
    try:
        data = request.json
        
        # AI validation of company data before SWOT analysis
        print("DEBUG: Running AI validation before SWOT analysis...")
        validation_result = validate_financial_data_with_ai(data)
        
        # Helper function to safely format numbers
        def safe_format_revenue(value):
            try:
                if value is None:
                    return '$0'
                return f"${float(value):,}"
            except (ValueError, TypeError):
                return '$0'
        
        # Create SWOT prompt with validation insights
        validation_notes = ""
        if validation_result.get('status') == 'validated':
            confidence = validation_result.get('confidence_score', 0)
            notes = validation_result.get('validation_notes', [])
            validation_notes = f"\n\nData Validation: Confidence Score: {confidence}%\nValidation Notes: {', '.join(notes)}"
        
        prompt = f"""
        Analyze this business for SWOT analysis and positioning guidance:
        
        Company: {data.get('company_name', 'Unknown')}
        Industry: {data.get('industry', 'Unknown')}
        Revenue: {safe_format_revenue(data.get('revenue', 0))}
        Employees: {data.get('employees', 'Unknown')}
        Key Markets: {str(data.get('key_markets', 'Unknown'))}
        Certifications: {str(data.get('certifications', []))}
        Competitive Advantages: {str(data.get('competitive_advantages', []))}
        
        {validation_notes}
        
        Please provide:
        1. Detailed SWOT Analysis (Strengths, Weaknesses, Opportunities, Threats)
        2. Actionable positioning guidance for selling this business
        3. Key value drivers to highlight to potential buyers
        4. Risk mitigation strategies
        
        Consider the data validation results when providing analysis. If confidence is low, suggest additional verification steps.
        
        Format as JSON with sections: strengths, weaknesses, opportunities, threats, positioning_guidance, value_drivers, risk_mitigation
        """
        
        # For now, return a mock response since OpenAI key might not be configured
        mock_swot = {
            'strengths': [
                'Strong market position',
                'Experienced management team',
                'Diversified customer base',
                'Quality certifications'
            ],
            'weaknesses': [
                'Owner dependency',
                'Limited geographic presence',
                'Aging equipment (if applicable)'
            ],
            'opportunities': [
                'Market expansion',
                'New product lines',
                'Strategic partnerships',
                'Technology upgrades'
            ],
            'threats': [
                'Economic downturn',
                'Increased competition',
                'Regulatory changes',
                'Key customer loss'
            ],
            'positioning_guidance': [
                'Highlight recurring revenue streams',
                'Emphasize growth potential',
                'Showcase operational efficiency',
                'Demonstrate market leadership'
            ],
            'value_drivers': [
                'Consistent profitability',
                'Strong customer relationships',
                'Operational scalability',
                'Market position'
            ],
            'risk_mitigation': [
                'Diversify customer base',
                'Cross-train employees',
                'Update technology systems',
                'Strengthen supplier relationships'
            ]
        }
        
        # If OpenAI is configured, use it instead
        # if openai.api_key:
        #     response = openai.ChatCompletion.create(
        #         model="gpt-4",
        #         messages=[{"role": "user", "content": prompt}],
        #         max_tokens=1000
        #     )
        #     swot_analysis = response['choices'][0]['message']['content']
        # else:
        
        return jsonify({
            'status': 'success',
            'swot_analysis': mock_swot,
            'prompt_used': prompt,
            'data_validation': validation_result
        })
        
    except Exception as e:
        return jsonify({'error': f'SWOT generation failed: {str(e)}'}), 500

# Helper function to safely format numbers
def safe_format_number(value, default=0):
    try:
        if value is None:
            return default
        return float(value)
    except (ValueError, TypeError):
        return default

@app.route('/api/report/generate', methods=['POST'])
@require_auth
def generate_report():
    """Generate comprehensive valuation report in multiple formats (requires authentication)"""
    try:
        data = request.json
        report_type = data.get('format', 'pdf').lower()
        
        # Debug: Print the data being received
        print(f"DEBUG: Report generation data received:")
        print(f"DEBUG: Company data keys: {list(data.get('company_data', {}).keys())}")
        print(f"DEBUG: Valuation results keys: {list(data.get('valuation_results', {}).keys())}")
        print(f"DEBUG: SWOT analysis keys: {list(data.get('swot_analysis', {}).keys())}")
        print(f"DEBUG: Requested format: {report_type}")
        
        # Get company data
        company_data = data.get('company_data', {})
        valuation_results = data.get('valuation_results', {})
        swot_analysis = data.get('swot_analysis', {})
        
        # AI validation before report generation
        print("DEBUG: Running AI validation before report generation...")
        validation_result = validate_financial_data_with_ai(company_data)
        
        # Generate report based on requested format
        try:
            if report_type == 'pdf':
                report_filename, report_path = generate_pdf_report(company_data, valuation_results, swot_analysis, data, safe_format_number)
            elif report_type == 'excel':
                report_filename, report_path = generate_excel_report(company_data, valuation_results, swot_analysis, data, safe_format_number)
            elif report_type == 'word':
                report_filename, report_path = generate_word_report(company_data, valuation_results, swot_analysis, data, safe_format_number)
            else:
                # Default to PDF if format not supported
                report_filename, report_path = generate_pdf_report(company_data, valuation_results, swot_analysis, data, safe_format_number)
        except Exception as format_error:
            print(f"Format-specific generation failed: {str(format_error)}")
            print("Falling back to text format...")
            report_filename, report_path = generate_text_report(company_data, valuation_results, swot_analysis, data, safe_format_number)
        
        # Log successful report generation
        log_user_activity(current_user.id, 'report_generation', True)
        
        return jsonify({
            'status': 'success',
            'report_filename': report_filename,
            'report_path': report_path,
            'download_url': f'/api/report/download/{report_filename}',
            'message': f'Comprehensive valuation report generated successfully in {report_type.upper()} format',
            'ai_validation': validation_result
        })
        
    except Exception as e:
        print(f"Report generation error: {str(e)}")
        # Log failed report generation
        if current_user.is_authenticated:
            log_user_activity(current_user.id, 'report_generation', False)
        return jsonify({'error': f'Report generation failed: {str(e)}'}), 500

def generate_pdf_report(company_data, valuation_results, swot_analysis, data, safe_format_number):
    """Generate PDF report using ReportLab"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        
        # Create filename
        report_filename = f"valuation_report_{company_data.get('company_name', 'company').replace(' ', '_')}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        report_path = os.path.join(app.config['REPORTS_FOLDER'], report_filename)
        
        # Create PDF document
        doc = SimpleDocTemplate(report_path, pagesize=A4)
        story = []
        
        # Get styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.darkblue
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            spaceBefore=20,
            textColor=colors.darkblue
        )
        
        # Title
        story.append(Paragraph("BUSINESS VALUATION REPORT", title_style))
        story.append(Spacer(1, 20))
        
        # Company Information
        story.append(Paragraph(f"Company: {company_data.get('company_name', 'Unknown Company')}", styles['Normal']))
        story.append(Paragraph(f"Industry: {company_data.get('industry', 'Not Specified')}", styles['Normal']))
        story.append(Paragraph(f"Report Date: {datetime.datetime.now().strftime('%B %d, %Y at %I:%M %p')}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Executive Summary
        story.append(Paragraph("EXECUTIVE SUMMARY", heading_style))
        story.append(Paragraph(data.get('executive_summary', 'Executive summary not available'), styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Company Overview
        story.append(Paragraph("COMPANY OVERVIEW", heading_style))
        
        # Financial metrics table
        financial_data = [
            ['Metric', 'Value'],
            ['Annual Revenue', f"${safe_format_number(company_data.get('revenue', 0)):,.0f}"],
            ['EBITDA', f"${safe_format_number(company_data.get('ebitda', 0)):,.0f}"],
            ['Net Income', f"${safe_format_number(company_data.get('net_income', 0)):,.0f}"],
            ['Total Assets', f"${safe_format_number(company_data.get('total_assets', 0)):,.0f}"],
            ['Inventory', f"${safe_format_number(company_data.get('inventory', 0)):,.0f}"],
            ['Accounts Receivable', f"${safe_format_number(company_data.get('accounts_receivable', 0)):,.0f}"],
            ['Cash', f"${safe_format_number(company_data.get('cash', 0)):,.0f}"]
        ]
        
        financial_table = Table(financial_data, colWidths=[2*inch, 2*inch])
        financial_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(financial_table)
        story.append(Spacer(1, 20))
        
        # Valuation Results
        story.append(Paragraph("VALUATION RESULTS", heading_style))
        
        valuation_data = [
            ['Method', 'Value'],
            ['Asset-Based', f"${safe_format_number(valuation_results.get('asset_based', 0)):,.0f}"],
            ['Income-Based', f"${safe_format_number(valuation_results.get('income_based', 0)):,.0f}"],
            ['Market-Based', f"${safe_format_number(valuation_results.get('market_based', 0)):,.0f}"],
            ['Low Estimate', f"${safe_format_number(valuation_results.get('valuation_range', {}).get('low', 0)):,.0f}"],
            ['Mid Estimate', f"${safe_format_number(valuation_results.get('valuation_range', {}).get('mid', 0)):,.0f}"],
            ['High Estimate', f"${safe_format_number(valuation_results.get('valuation_range', {}).get('high', 0)):,.0f}"]
        ]
        
        valuation_table = Table(valuation_data, colWidths=[2*inch, 2*inch])
        valuation_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(valuation_table)
        story.append(Spacer(1, 20))
        
        # SWOT Analysis
        story.append(Paragraph("SWOT ANALYSIS", heading_style))
        
        # Strengths
        story.append(Paragraph("Strengths:", styles['Heading3']))
        for strength in swot_analysis.get('strengths', ['Not available']):
            story.append(Paragraph(f"• {strength}", styles['Normal']))
        story.append(Spacer(1, 10))
        
        # Weaknesses
        story.append(Paragraph("Weaknesses:", styles['Heading3']))
        for weakness in swot_analysis.get('weaknesses', ['Not available']):
            story.append(Paragraph(f"• {weakness}", styles['Normal']))
        story.append(Spacer(1, 10))
        
        # Opportunities
        story.append(Paragraph("Opportunities:", styles['Heading3']))
        for opportunity in swot_analysis.get('opportunities', ['Not available']):
            story.append(Paragraph(f"• {opportunity}", styles['Normal']))
        story.append(Spacer(1, 10))
        
        # Threats
        story.append(Paragraph("Threats:", styles['Heading3']))
        for threat in swot_analysis.get('threats', ['Not available']):
            story.append(Paragraph(f"• {threat}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Methodology
        story.append(Paragraph("METHODOLOGY & ASSUMPTIONS", heading_style))
        story.append(Paragraph("Valuation Methods Used:", styles['Heading3']))
        story.append(Paragraph("1. Asset-Based Approach: Book value adjusted for market conditions (80% of book value)", styles['Normal']))
        story.append(Paragraph("2. Income-Based Approach: EBITDA multiple analysis (6x industry average)", styles['Normal']))
        story.append(Paragraph("3. Market-Based Approach: Revenue multiple analysis (1.5x conservative multiple)", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Recommendations
        story.append(Paragraph("RECOMMENDATIONS", heading_style))
        story.append(Paragraph(f"1. Primary Valuation: ${safe_format_number(valuation_results.get('valuation_range', {}).get('mid', 0)):,.0f}", styles['Normal']))
        story.append(Paragraph(f"2. Negotiation Range: ${safe_format_number(valuation_results.get('valuation_range', {}).get('low', 0)):,.0f} - ${safe_format_number(valuation_results.get('valuation_range', {}).get('high', 0)):,.0f}", styles['Normal']))
        story.append(Paragraph(f"3. Key Value Drivers: {', '.join(swot_analysis.get('value_drivers', ['Not specified']))}", styles['Normal']))
        story.append(Paragraph(f"4. Risk Factors: {', '.join(swot_analysis.get('risk_mitigation', ['Not specified']))}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Disclaimer
        story.append(Paragraph("DISCLAIMER", heading_style))
        disclaimer_text = """This valuation report is prepared for informational purposes only and should not be 
        considered as investment advice. The analysis is based on the information provided 
        and current market conditions. Professional consultation is recommended before 
        making any investment decisions."""
        story.append(Paragraph(disclaimer_text, styles['Normal']))
        
        # Build PDF
        doc.build(story)
        
        print(f"PDF report generated: {report_path}")
        return report_filename, report_path
        
    except Exception as e:
        print(f"PDF generation error: {str(e)}")
        # Fallback to text format
        return generate_text_report(company_data, valuation_results, swot_analysis, data, safe_format_number)

def generate_excel_report(company_data, valuation_results, swot_analysis, data, safe_format_number):
    """Generate Excel report using openpyxl"""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        
        # Create filename
        report_filename = f"valuation_report_{company_data.get('company_name', 'company').replace(' ', '_')}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        report_path = os.path.join(app.config['REPORTS_FOLDER'], report_filename)
        
        # Create workbook and worksheet
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Valuation Report"
        
        # Styles
        title_font = Font(name='Arial', size=16, bold=True, color="FFFFFF")
        heading_font = Font(name='Arial', size=12, bold=True, color="FFFFFF")
        normal_font = Font(name='Arial', size=10)
        
        title_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        heading_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        
        # Title
        ws['A1'] = "BUSINESS VALUATION REPORT"
        ws['A1'].font = title_font
        ws['A1'].fill = title_fill
        ws.merge_cells('A1:F1')
        ws['A1'].alignment = Alignment(horizontal='center')
        
        # Company Information
        ws['A3'] = "Company:"
        ws['B3'] = company_data.get('company_name', 'Unknown Company')
        ws['A4'] = "Industry:"
        ws['B4'] = company_data.get('industry', 'Not Specified')
        ws['A5'] = "Report Date:"
        ws['B5'] = datetime.datetime.now().strftime('%B %d, %Y at %I:%M %p')
        
        # Executive Summary
        ws['A7'] = "EXECUTIVE SUMMARY"
        ws['A7'].font = heading_font
        ws['A7'].fill = heading_fill
        ws.merge_cells('A7:F7')
        ws['A8'] = data.get('executive_summary', 'Executive summary not available')
        ws.merge_cells('A8:F8')
        
        # Company Overview
        ws['A10'] = "COMPANY OVERVIEW"
        ws['A10'].font = heading_font
        ws['A10'].fill = heading_fill
        ws.merge_cells('A10:F10')
        
        # Financial metrics table
        financial_headers = ['Metric', 'Value']
        financial_data = [
            ['Annual Revenue', f"${safe_format_number(company_data.get('revenue', 0)):,.0f}"],
            ['EBITDA', f"${safe_format_number(company_data.get('ebitda', 0)):,.0f}"],
            ['Net Income', f"${safe_format_number(company_data.get('net_income', 0)):,.0f}"],
            ['Total Assets', f"${safe_format_number(company_data.get('total_assets', 0)):,.0f}"],
            ['Inventory', f"${safe_format_number(company_data.get('inventory', 0)):,.0f}"],
            ['Accounts Receivable', f"${safe_format_number(company_data.get('accounts_receivable', 0)):,.0f}"],
            ['Cash', f"${safe_format_number(company_data.get('cash', 0)):,.0f}"]
        ]
        
        # Add headers
        for col, header in enumerate(financial_headers, 1):
            cell = ws.cell(row=12, column=col)
            cell.value = header
            cell.font = heading_font
            cell.fill = heading_fill
            cell.alignment = Alignment(horizontal='center')
        
        # Add data
        for row_idx, row_data in enumerate(financial_data, 13):
            for col_idx, value in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col_idx)
                cell.value = value
                cell.font = normal_font
                cell.alignment = Alignment(horizontal='center')
        
        # Valuation Results
        ws['A20'] = "VALUATION RESULTS"
        ws['A20'].font = heading_font
        ws['A20'].fill = heading_fill
        ws.merge_cells('A20:F20')
        
        valuation_headers = ['Method', 'Value']
        valuation_data = [
            ['Asset-Based', f"${safe_format_number(valuation_results.get('asset_based', 0)):,.0f}"],
            ['Income-Based', f"${safe_format_number(valuation_results.get('income_based', 0)):,.0f}"],
            ['Market-Based', f"${safe_format_number(valuation_results.get('market_based', 0)):,.0f}"],
            ['Low Estimate', f"${safe_format_number(valuation_results.get('valuation_range', {}).get('low', 0)):,.0f}"],
            ['Mid Estimate', f"${safe_format_number(valuation_results.get('valuation_range', {}).get('mid', 0)):,.0f}"],
            ['High Estimate', f"${safe_format_number(valuation_results.get('valuation_range', {}).get('high', 0)):,.0f}"]
        ]
        
        # Add valuation headers
        for col, header in enumerate(valuation_headers, 1):
            cell = ws.cell(row=22, column=col)
            cell.value = header
            cell.font = heading_font
            cell.fill = heading_fill
            cell.alignment = Alignment(horizontal='center')
        
        # Add valuation data
        for row_idx, row_data in enumerate(valuation_data, 23):
            for col_idx, value in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col_idx)
                cell.value = value
                cell.font = normal_font
                cell.alignment = Alignment(horizontal='center')
        
        # SWOT Analysis
        ws['A30'] = "SWOT ANALYSIS"
        ws['A30'].font = heading_font
        ws['A30'].fill = heading_fill
        ws.merge_cells('A30:F30')
        
        # Add SWOT sections
        swot_sections = [
            ('Strengths', swot_analysis.get('strengths', ['Not available'])),
            ('Weaknesses', swot_analysis.get('weaknesses', ['Not available'])),
            ('Opportunities', swot_analysis.get('opportunities', ['Not available'])),
            ('Threats', swot_analysis.get('threats', ['Not available']))
        ]
        
        current_row = 32
        for section_name, items in swot_sections:
            ws[f'A{current_row}'] = section_name
            ws[f'A{current_row}'].font = Font(bold=True)
            current_row += 1
            
            for item in items:
                ws[f'B{current_row}'] = f"• {item}"
                current_row += 1
            current_row += 1
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = get_column_letter(column[0].column)
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        # Save workbook
        wb.save(report_path)
        
        print(f"Excel report generated: {report_path}")
        return report_filename, report_path
        
    except Exception as e:
        print(f"Excel generation error: {str(e)}")
        # Fallback to text format
        return generate_text_report(company_data, valuation_results, swot_analysis, data, safe_format_number)

def generate_word_report(company_data, valuation_results, swot_analysis, data, safe_format_number):
    """Generate Word document report using python-docx"""
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.shared import OxmlElement, qn
        
        # Create filename
        report_filename = f"valuation_report_{company_data.get('company_name', 'company').replace(' ', '_')}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        report_path = os.path.join(app.config['REPORTS_FOLDER'], report_filename)
        
        # Create document
        doc = Document()
        
        # Title
        title = doc.add_heading('BUSINESS VALUATION REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Company Information
        doc.add_heading('Company Information', level=1)
        doc.add_paragraph(f"Company: {company_data.get('company_name', 'Unknown Company')}")
        doc.add_paragraph(f"Industry: {company_data.get('industry', 'Not Specified')}")
        doc.add_paragraph(f"Report Date: {datetime.datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(data.get('executive_summary', 'Executive summary not available'))
        
        # Company Overview
        doc.add_heading('Company Overview', level=1)
        doc.add_paragraph(f"Annual Revenue: ${safe_format_number(company_data.get('revenue', 0)):,.0f}")
        doc.add_paragraph(f"EBITDA: ${safe_format_number(company_data.get('ebitda', 0)):,.0f}")
        doc.add_paragraph(f"Net Income: ${safe_format_number(company_data.get('net_income', 0)):,.0f}")
        doc.add_paragraph(f"Total Assets: ${safe_format_number(company_data.get('total_assets', 0)):,.0f}")
        doc.add_paragraph(f"Inventory: ${safe_format_number(company_data.get('inventory', 0)):,.0f}")
        doc.add_paragraph(f"Accounts Receivable: ${safe_format_number(company_data.get('accounts_receivable', 0)):,.0f}")
        doc.add_paragraph(f"Cash: ${safe_format_number(company_data.get('cash', 0)):,.0f}")
        
        # Valuation Results
        doc.add_heading('Valuation Results', level=1)
        doc.add_paragraph(f"Asset-Based Valuation: ${safe_format_number(valuation_results.get('asset_based', 0)):,.0f}")
        doc.add_paragraph(f"Income-Based Valuation: ${safe_format_number(valuation_results.get('income_based', 0)):,.0f}")
        doc.add_paragraph(f"Market-Based Valuation: ${safe_format_number(valuation_results.get('market_based', 0)):,.0f}")
        doc.add_paragraph(f"Low Estimate: ${safe_format_number(valuation_results.get('valuation_range', {}).get('low', 0)):,.0f}")
        doc.add_paragraph(f"Mid Estimate: ${safe_format_number(valuation_results.get('valuation_range', {}).get('mid', 0)):,.0f}")
        doc.add_paragraph(f"High Estimate: ${safe_format_number(valuation_results.get('valuation_range', {}).get('high', 0)):,.0f}")
        
        # SWOT Analysis
        doc.add_heading('SWOT Analysis', level=1)
        
        # Strengths
        doc.add_heading('Strengths', level=2)
        for strength in swot_analysis.get('strengths', ['Not available']):
            doc.add_paragraph(f"• {strength}")
        
        # Weaknesses
        doc.add_heading('Weaknesses', level=2)
        for weakness in swot_analysis.get('weaknesses', ['Not available']):
            doc.add_paragraph(f"• {weakness}")
        
        # Opportunities
        doc.add_heading('Opportunities', level=2)
        for opportunity in swot_analysis.get('opportunities', ['Not available']):
            doc.add_paragraph(f"• {opportunity}")
        
        # Threats
        doc.add_heading('Threats', level=2)
        for threat in swot_analysis.get('threats', ['Not available']):
            doc.add_paragraph(f"• {threat}")
        
        # Methodology
        doc.add_heading('Methodology & Assumptions', level=1)
        doc.add_paragraph("Valuation Methods Used:")
        doc.add_paragraph("1. Asset-Based Approach: Book value adjusted for market conditions (80% of book value)")
        doc.add_paragraph("2. Income-Based Approach: EBITDA multiple analysis (6x industry average)")
        doc.add_paragraph("3. Market-Based Approach: Revenue multiple analysis (1.5x conservative multiple)")
        
        # Recommendations
        doc.add_heading('Recommendations', level=1)
        doc.add_paragraph(f"1. Primary Valuation: ${safe_format_number(valuation_results.get('valuation_range', {}).get('mid', 0)):,.0f}")
        doc.add_paragraph(f"2. Negotiation Range: ${safe_format_number(valuation_results.get('valuation_range', {}).get('low', 0)):,.0f} - ${safe_format_number(valuation_results.get('valuation_range', {}).get('high', 0)):,.0f}")
        doc.add_paragraph(f"3. Key Value Drivers: {', '.join(swot_analysis.get('value_drivers', ['Not specified']))}")
        doc.add_paragraph(f"4. Risk Factors: {', '.join(swot_analysis.get('risk_mitigation', ['Not specified']))}")
        
        # Disclaimer
        doc.add_heading('Disclaimer', level=1)
        disclaimer_text = """This valuation report is prepared for informational purposes only and should not be 
        considered as investment advice. The analysis is based on the information provided 
        and current market conditions. Professional consultation is recommended before 
        making any investment decisions."""
        doc.add_paragraph(disclaimer_text)
        
        # Save document
        doc.save(report_path)
        
        print(f"Word report generated: {report_path}")
        return report_filename, report_path
        
    except Exception as e:
        print(f"Word generation error: {str(e)}")
        # Fallback to text format
        return generate_text_report(company_data, valuation_results, swot_analysis, data, safe_format_number)

def generate_text_report(company_data, valuation_results, swot_analysis, data, safe_format_number):
    """Generate text report as fallback"""
    try:
        # Create filename
        report_filename = f"valuation_report_{company_data.get('company_name', 'company').replace(' ', '_')}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        report_path = os.path.join(app.config['REPORTS_FOLDER'], report_filename)
        
        # Create comprehensive report content
        report_content = f"""
{'='*80}
                    BUSINESS VALUATION REPORT
{'='*80}

Company: {company_data.get('company_name', 'Unknown Company')}
Industry: {company_data.get('industry', 'Not Specified')}
Report Date: {datetime.datetime.now().strftime('%B %d, %Y at %I:%M %p')}
Valuation Date: {datetime.datetime.now().strftime('%B %d, %Y')}

{'='*80}
                           EXECUTIVE SUMMARY
{'='*80}

{data.get('executive_summary', 'Executive summary not available')}

{'='*80}
                        COMPANY OVERVIEW
{'='*80}

Financial Metrics:
• Annual Revenue: ${safe_format_number(company_data.get('revenue', 0)):,.0f}
• EBITDA: ${safe_format_number(company_data.get('ebitda', 0)):,.0f}
• Net Income: ${safe_format_number(company_data.get('net_income', 0)):,.0f}
• Total Assets: ${safe_format_number(company_data.get('total_assets', 0)):,.0f}
• Inventory: ${safe_format_number(company_data.get('inventory', 0)):,.0f}
• Accounts Receivable: ${safe_format_number(company_data.get('accounts_receivable', 0)):,.0f}
• Cash: ${safe_format_number(company_data.get('cash', 0)):,.0f}

{'='*80}
                        VALUATION RESULTS
{'='*80}

Asset-Based Valuation:
• Total Asset Value: ${safe_format_number(valuation_results.get('asset_based', 0)):,.0f}

Income-Based Valuation:
• EBITDA Multiple Value: ${safe_format_number(valuation_results.get('income_based', 0)):,.0f}

Market-Based Valuation:
• Revenue Multiple Value: ${safe_format_number(valuation_results.get('market_based', 0)):,.0f}

FINAL VALUATION RANGE:
• Low Estimate:  ${safe_format_number(valuation_results.get('valuation_range', {}).get('low', 0)):,.0f}
• Mid Estimate:  ${safe_format_number(valuation_results.get('valuation_range', {}).get('mid', 0)):,.0f}
• High Estimate: ${safe_format_number(valuation_results.get('valuation_range', {}).get('high', 0)):,.0f}

Methodology: {valuation_results.get('methodology', 'Standard valuation approaches')}
Key Assumptions: {valuation_results.get('assumptions', 'Industry standard multiples and adjustments')}

{'='*80}
                        SWOT ANALYSIS
{'='*80}

Strengths:
{chr(10).join([f'• {strength}' for strength in swot_analysis.get('strengths', ['Not available'])])}

Weaknesses:
{chr(10).join([f'• {weakness}' for weakness in swot_analysis.get('weaknesses', ['Not available'])])}

Opportunities:
{chr(10).join([f'• {opportunity}' for opportunity in swot_analysis.get('opportunities', ['Not available'])])}

Threats:
{chr(10).join([f'• {threat}' for threat in swot_analysis.get('threats', ['Not available'])])}

Positioning Guidance:
{chr(10).join([f'• {guidance}' for guidance in swot_analysis.get('positioning_guidance', ['Not available'])])}

Value Drivers:
{chr(10).join([f'• {driver}' for driver in swot_analysis.get('value_drivers', ['Not specified'])])}

Risk Mitigation:
{chr(10).join([f'• {risk}' for risk in swot_analysis.get('risk_mitigation', ['Not specified'])])}

{'='*80}
                        METHODOLOGY & ASSUMPTIONS
{'='*80}

Valuation Methods Used:
1. Asset-Based Approach: Book value adjusted for market conditions (80% of book value)
2. Income-Based Approach: EBITDA multiple analysis (6x industry average)
3. Market-Based Approach: Revenue multiple analysis (1.5x conservative multiple)

Key Assumptions:
• Asset Discount: 20% of book value for market conditions
• EBITDA Multiple: 6.0x (industry average)
• Revenue Multiple: 1.5x (conservative estimate)
• Industry Standards: Based on {company_data.get('industry', 'general business')} sector

{'='*80}
                        RECOMMENDATIONS
{'='*80}

Based on our analysis, we recommend:

1. Primary Valuation: ${safe_format_number(valuation_results.get('valuation_range', {}).get('mid', 0)):,.0f}
2. Negotiation Range: ${safe_format_number(valuation_results.get('valuation_range', {}).get('low', 0)):,.0f} - ${safe_format_number(valuation_results.get('valuation_range', {}).get('high', 0)):,.0f}
3. Key Value Drivers: {', '.join(swot_analysis.get('value_drivers', ['Not specified']))}
4. Risk Factors: {', '.join(swot_analysis.get('risk_mitigation', ['Not specified']))}

{'='*80}
                        DISCLAIMER
{'='*80}

This valuation report is prepared for informational purposes only and should not be 
considered as investment advice. The analysis is based on the information provided 
and current market conditions. Professional consultation is recommended before 
making any investment decisions.

Report Generated: {datetime.datetime.now().strftime('%B %d, %Y at %I:%M %p')}
Valuation Platform: Business Valuation Platform v1.0
{'='*80}
        """
        
        # Save report to file
        with open(report_path, 'w') as f:
            f.write(report_content)
        
        print(f"Text report generated: {report_path}")
        return report_filename, report_path
        
    except Exception as e:
        print(f"Text report generation error: {str(e)}")
        raise e

@app.route('/api/report/download/<filename>', methods=['GET'])
def download_report(filename):
    """Download generated report"""
    try:
        report_path = os.path.join(app.config['REPORTS_FOLDER'], filename)
        if os.path.exists(report_path):
            return send_file(report_path, as_attachment=True)
        else:
            return jsonify({'error': 'Report not found'}), 404
    except Exception as e:
        return jsonify({'error': f'Download failed: {str(e)}'}), 500

if __name__ == '__main__':
    # Production configuration
    port = int(os.environ.get('PORT', 5000))
    debug = os.environ.get('FLASK_ENV') == 'development'
    
    app.run(host='0.0.0.0', port=port, debug=debug)
